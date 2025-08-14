from docx import Document
import io
from flask import Flask, request, render_template_string, redirect, send_file
from datetime import timedelta
from dateutil.relativedelta import relativedelta

app = Flask(__name__)
landlords = {
    "Mr. Rajesh Khanna": {
        "Landlord Address": "A-15, West Patel Nagar",
        "Property Address": "xyz baljeet nagar"
    },
    "Mr. Pawan Kumar": {
        "Landlord Address": "A-16, West Patel Nagar",
        "Property Address": "xyz baljeet nagar"
    },
    "Mr. Harsh": {
        "Landlord Address": "A-17, West Patel Nagar",
        "Property Address": "xyz baljeet nagar"
    },
    "Mr. Kultaj Khanna": {
        "Landlord Address": "A-18, West Patel Nagar",
        "Property Address": "xyz baljeet nagar"
    },
    "Ms Preeti khanna": {
        "Landlord Address": "A-19, West Patel Nagar",
        "Property Address": "xyz baljeet nagar"
     }
}

form_template = '''
<style>
    body {
        font-family: Arial, sans-serif;
        background-color: #f4f6f8;
        display: flex;
        justify-content: center;
        align-items: center;
        min-height: 100vh;
        margin: 0;
    }

    .form-container {
        background-color: #ffffff;
        padding: 10px 15px;
        border-radius: 10px;
        box-shadow: 0 4px 10px rgba(0, 0, 0, 0.1);
        width: 90%;
        max-width: 450px;
        margin: 14px auto;
    }

    h2 {
        text-align: center;
        color: #333;
        margin-bottom: 6px;
        font-size: 14px;
    }

    label {
        display: block;
        margin-top: 6px;
        margin-bottom: 2px;
        font-weight: bold;
        font-size: 10px;
        color: #333;
    }

    input[type="text"],
    input[type="number"],
    input[type="date"],
    select {
        width: 100%;
        padding:1px 1px;
        border-radius: 1px;
        border: 1px solid #ccc;
        box-sizing: border-box;
        margin-bottom: 1px;
    }

    button {
        width: 100%;
        background-color: #007BFF;
        color: white;
        padding: 4px;
        margin-top: 4px;
        border: none;
        border-radius: 3px;
        cursor: pointer;
        font-size: 10px;
    }

    button:hover {
        background-color: #0056b3;
    }

        .spacer {
            height: 30px;
    }
</style>


<div class = "form-container">
<form method="post">

    Landlord Name:
    <select name="landlord_name" onchange="this.form.submit()">
    <option value="">--Select--</option>
    {% for name in landlords %}
        <option value="{{ name }}" {% if landlord_name == name %}selected{% endif %} required>{{ name }}</option>
    {% endfor %}
     </select>
    <br><br>  

    Tenant Name: <input type="text" name="tenant_name" value="{{ tenant_name }}" required><br><br>
    Landlord Address: <input type="text" name="landlord_address" value="{{ landlord_address }}" readonly><br><br>

    
    Tenant Address: <input type="text" name="tenant_address" value="{{ tenant_address }}" required><br><br>
    Rent Amount: <input type="number" name="rent_amount" value="{{ rent_amount }}" required><br><br>
    Security Amount: <input type="number" name="security_amount" value="{{ security_amount }}" required><br><br>
    Tenancy Duration(months): <input type="number" name="tenancy_duration" value="{{ tenancy_duration }}" required><br><br>
    Tenancy Start Date: <input type="date" name="tenancy_start" value="{{ tenancy_start }}" required><br><br>
    Rent Start Date: <input type="date" name="rent_start" value="{{ rent_start }}" required><br><br>

    Property Address: <input type="text" name="property_address" value="{{ property_address }}" readonly><br><br>

    <button type="submit" name="action" value="preview">Preview</button>
    <button type="submit" name="action" value="download">Download</button>
</form>
</div>
'''

preview_template = '''
<h2>Preview</h2>
<h3>Preview Agreement</h3>
<pre>{{agreement_text}}</pre>
<form method="post">
    <input type="hidden" name="landlord_name" value="{{ landlord_name }}">
    <input type="hidden" name="tenant_name" value="{{ tenant_name }}">
    <input type="hidden" name="landlord_address" value="{{ landlord_address }}">
    <input type="hidden" name="tenant_address" value="{{ tenant_address}}">
    <input type="hidden" name="rent_amount" value="{{ rent_amount }}">
    <input type="hidden" name="security_amount" value="{{ security_amount }}">
    <input type="hidden" name="tenancy_duration" value="{{ tenancy_duration }}">
    <input type="hidden" name="tenancy_start" value="{{ tenancy_start }}">
    <input type="hidden" name="rent_start" value="{{ rent_start }}">
    <input type="hidden" name="property_address" value="{{ property_address }}">
    <p><b>Tenancy End Date:</b> {{ tenancy_end_str }}</p>


    
    <button type="submit" name="action" value="download">Download</button>
    <button type="submit" name="action" value="go_back">Go Back</button>
</form>
'''
from datetime import datetime
@app.route('/', methods=['GET', 'POST'])
def main():  
# default form values
    form_data= {
        'landlord_name' : '',
        'landlord_address' : '',
        'property_address' : '',
        'tenant_name' : '',
        'tenant_address' : '',
        'rent_amount' : '',
        'security_amount' : '',
        'tenancy_duration' : '',
        'tenancy_start' : '',
        'rent_start' : ''

    }

    if request.method == 'POST':
        action = request.form.get("action", "")
        landlord_name = request.form.get('landlord_name', '')

        action = request.form.get('action')
        form_data=request.form.to_dict()

        #autofill landlord details

        if landlord_name in landlords:
           landlord_address = landlords[landlord_name]["Landlord Address"]
           property_address = landlords[landlord_name]["Property Address"]

           
           form_data['landlord_address'] = landlord_address
           form_data['property_address'] = property_address
        # Extract values
           landlord_name = request.form.get('landlord_name', '')
           tenant_name = form_data.get('tenant_name', '')
           landlord_address = form_data.get('landlord_address', '')
           tenant_address = form_data.get('tenant_address', '')
           rent_amount = form_data.get('rent_amount', '')
           security_amount = form_data.get('security_amount', '')
           tenancy_duration = form_data.get('tenancy_duration', '')
           tenancy_start = form_data.get('tenancy_start', '')
           rent_start = form_data.get('rent_start', '')
           property_address = form_data.get('property_address', '')

        # Format date

        # Inside your function:
         
           if action == 'go_back':
            return render_template_string(form_template, landlords=landlords, **form_data)

           if action == 'preview' or action == 'download':
               try:
                
                   start_date = datetime.strptime(tenancy_start, '%Y-%m-%d')
                   end_date = start_date + relativedelta(months=int(tenancy_duration))
                   tenancy_end_str = end_date.strftime('%d-%m-%Y')

                   formatted_date = datetime.now().strftime('%d-%m-%Y')

        #agreement text:

                   agreement_text=f"""This RENT AGREEMENT is made at Delhi on { formatted_date }
                          
    Between

    Shri {landlord_name }, R/o { landlord_address } (hereinafter called the First Party/Landlord/Landlady)

    And

    { tenant_name }, R/o { tenant_address } (hereinafter called the Second Party/Tenant).

    The expressions of both the parties shall include their legal heirs, nominees, successors, administrators, etc.

    WHEREAS the First Party is the absolute owner and in possession of property located at { property_address } and has agreed to let out the same to the Second Party for a period of {tenancy_duration} months on a monthly rent of Rs { rent_amount }/- upon the request of the Second Party.
  
    NOW THIS RENT AGREEMENT WITNESSETH AS UNDER:

    1.Tenure: The tenancy period is { tenancy_duration } months from { tenancy_start } to {tenancy_end_str}. Any extension will require mutual consent and will involve a minimum 10% increase in rent.

    2.Rent Payment: The monthly rent of Rs { rent_amount }/- shall be paid in advance by the 9th of each month.

    3.Security Deposit: An interest-free security deposit of Rs { security_amount }/- has been paid by the Second Party, refundable upon vacating the premises, subject to conditions.

    4.	Utility Payments: The Second Party shall pay electricity and water bill to Government and share the receipt with first party.  
    5.	Non-Payment of Dues: Failure to pay rent or utility bills on time may lead to termination of the agreement, forfeiture of the security deposit, and legal action.
    6.	Agreement Custody: The original agreement shall remain with the First Party.
    7.	Alterations: No structural changes are allowed without prior written consent.
    8.	Lock-In Period: A six-month lock-in period applies. If the tenants vacate before this period ends, the security deposit will be forfeited. The First Party may terminate the agreement anytime with one month's notice.
    9.	Termination Notice: After the lock-in period, either party may terminate the agreement with one month's written notice.
    10.	Compliance with Apartment Rules: Tenants must adhere to building rules and avoid disturbances, or risk immediate termination.
    11.	Condition on Handover: Premises must be returned in a clean and undamaged condition. Costs for repairs or cleaning due to negligence will be deducted from the security deposit.  
    12.	Repairs and Maintenance: Minor repairs are the tenant's responsibility. The electric motor is guaranteed for the first month only; subsequent repairs are the tenant's responsibility.
    13.	Renewal: Renewal is subject to mutual agreement and a minimum 10% rent increase.
    14.	Overstay Penalty: Post-agreement occupancy will incur a penalty of Rs 1,000/- per day.
    15.	Permitted Use: The premises shall be used solely for residential purposes. Illegal or immoral activities will result in immediate eviction and security deposit forfeiture.
    16.	Subletting: Subletting is strictly prohibited. Unauthorized occupants will lead to agreement termination and security deposit forfeiture.
    17.	Breach of Terms: Violation of any terms may result in immediate termination, eviction, and forfeiture of the security deposit.
    18.	Furnishings: The premises include the following items as per Annexure I.
    19.	Wall Damage Prohibition: Pasting or nailing anything on walls without prior permission is strictly prohibited. Even if permission is granted, any damage to the wall paint will result in deduction of painting charges from the security deposit.
    In WITNESS WHEREOF this agreement has been signed on the date mentioned above.
    WITNESSES:
    1.	Rajesh Khanna S/o Late Shri Dharambir Khanna R/O A- 15, West Patel Nagar, New Delhi-110008
    2.	Manjeet Kaur W/O Balwinder Singh T 510-E5, Hill Marg, Baljeet Nagar, New Delhi-110008
    SIGNATORIES:
    First Party: ____________________ {landlord_name}
     Second Party:  ___________________ {tenant_name}  
                   
Annexure I - Inventory:
1.	Fan - 2
2.	Light - 7
3.	Bed - 2
4.	Table - 1
5.	Chair – 1
6.	Rack-1
7.	Air Conditioner
8.	Fridge
9.	Almirah



SIGNATORIES:

First Party: ____________________ ({landlord_name })

Second Party: ____________________ ({ tenant_name })"""
        
        
        #
                   if action == 'preview':
                    return render_template_string(preview_template, agreement_text=agreement_text,
                                           form_data=form_data, landlord_name=landlord_name,
                                           tenant_name=tenant_name,landlord_address=landlord_address,
                                           tenant_address=tenant_address,rent_amount=rent_amount,
                                           security_amount=security_amount,rent_start=rent_start,
                                           tenancy_duration=tenancy_duration,tenancy_start=tenancy_start,
                                           tenancy_end_str=tenancy_end_str, landlords=landlords )
                   elif action == 'go_back':
            
                    return render_template_string(form_template, **form_data,landlords=landlords)
           
                   elif action == 'download':
                    
                        doc = Document()
                        doc.add_heading("Rent Agreement", 0)
                        doc.add_paragraph(agreement_text)

                        file_stream = io.BytesIO()
                        doc.save(file_stream)
                        file_stream.seek(0)

                        return send_file(file_stream, as_attachment=True, download_name='Rent_Agreement.docx')

               except Exception as e:
                       return f"<h2>Error: {str(e)}</h2>"


    return render_template_string(form_template ,**form_data ,landlords=landlords)


               
if __name__ == '__main__':
                 
                 app.run(host='192.168.87.8', port=5000, debug=True)
